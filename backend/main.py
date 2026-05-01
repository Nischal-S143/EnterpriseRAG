"""
Pagani Zonda R – Enterprise Intelligence API
FastAPI backend with RAG, JWT auth, rate limiting, CORS, logging,
database persistence, security middleware, health monitoring,
enterprise RBAC, analytics, audit, document management, and WebSocket support.
"""

import os
import time
import logging
from datetime import datetime, timezone
from contextlib import asynccontextmanager
import asyncio
import hashlib
from typing import Optional

from fastapi import (
    FastAPI, Depends, HTTPException, Request, status,
    APIRouter, WebSocket, WebSocketDisconnect,
    Query as QueryParam,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, PlainTextResponse, FileResponse
from sqlalchemy.orm import Session
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import bleach
from pydantic import BaseModel, Field, field_validator
from dotenv import load_dotenv

from auth import (
    UserRegister, UserLogin, TokenResponse, RefreshRequest,
    ChatRequest, ChatResponse, UserInfo, register_user, authenticate_user, refresh_access_token,
    get_current_user, users_db, require_permission, verify_admin_key,
    ROLE_PERMISSIONS, VALID_ROLES,
)
from vector_store import vector_store
from rag_pipeline import (
    generate_response, 
    generate_response_stream,
    agentic_router,
    _get_history,
)
from logging_config import setup_logging, log_event
from database import check_db_connection, get_db
from models import Document, DocumentVersion
from middleware import SecurityHeadersMiddleware, RequestSizeLimitMiddleware, RequestTracingMiddleware
from error_handlers import register_error_handlers
from audit import audit, get_audit_logs, get_login_attempts
from analytics import (
    get_user_engagement_metrics, get_query_success_rates,
    get_ai_performance_metrics, get_system_health,
    export_analytics_csv, get_analytics_summary, Strategist, set_server_start_time,
    track_session_start, track_session_end,
)
from websocket_manager import ws_manager
from cache import query_cache
from evaluator import Evaluator, IRMetrics
from stress_tester import StressTester
from auth import review_queue as auth_review_queue
from sse_manager import sse_manager

# Load environment variables from the same directory as this file
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

# ── Structured Logging (replaces basicConfig) ──
setup_logging(level=os.getenv("LOG_LEVEL", "INFO"))
logger = logging.getLogger("pagani.api")

# ── Rate Limiter ──
limiter = Limiter(key_func=get_remote_address)

# ── Server Start Time (for uptime tracking) ──
SERVER_START_TIME = None

PIPELINE_STATUS = {
    "data_sources": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "restructuring": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "chunking": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "metadata": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "planner": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "tool_execution": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "router": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "multi_agent": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "agent_1": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "agent_2": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "agent_3": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "human_validation": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "evaluation": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
    "database": {"status": "idle", "last_run_ms": 0, "last_run_at": None},
}

def update_pipeline_node(node: str, status: str, duration: int = 0):
    if node in PIPELINE_STATUS:
        PIPELINE_STATUS[node]["status"] = status
        if duration > 0:
            PIPELINE_STATUS[node]["last_run_ms"] = duration
        PIPELINE_STATUS[node]["last_run_at"] = datetime.now(timezone.utc).isoformat()


# ── Lifespan ──
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize startup tasks in the background while uvicorn starts."""
    global SERVER_START_TIME
    SERVER_START_TIME = datetime.now(timezone.utc)
    set_server_start_time(SERVER_START_TIME)

    logger.info("═" * 60)
    logger.info("  PAGANI ZONDA R — Enterprise Intelligence API (Loading...)")
    logger.info("═" * 60)

    async def _background_init():
        """Handles heavy-duty initialization without blocking the event loop."""
        import time
        # Initialize database
        try:
            update_pipeline_node("database", "running")
            t0 = time.time()
            from database import init_db
            init_db()
            
            from auth import _load_users_from_db
            _load_users_from_db()
            dur = int((time.time() - t0) * 1000)
            update_pipeline_node("database", "done", dur)
            logger.info("Background: Database initialized and users loaded.")
        except Exception as e:
            update_pipeline_node("database", "error")
            logger.error(f"Background: Database initialization failed: {e}")

        # Initialize vector store
        try:
            update_pipeline_node("data_sources", "running")
            update_pipeline_node("chunking", "running")
            update_pipeline_node("metadata", "running")
            t0 = time.time()
            vector_store.initialize()
            dur = int((time.time() - t0) * 1000)
            update_pipeline_node("data_sources", "done", dur)
            update_pipeline_node("chunking", "done", max(10, dur // 3))
            update_pipeline_node("metadata", "done", max(10, dur // 3))
            update_pipeline_node("restructuring", "done", max(10, dur // 4))
            logger.info("Background: Vector store initialized.")
        except Exception as e:
            update_pipeline_node("metadata", "error")
            logger.error(f"Background: Vector store initialization failed: {e}")

        # Initialize AI Strategist
        try:
            strategist = Strategist()
            await strategist.start_background_task()
            logger.info("Background: AI Strategist started.")
        except Exception as e:
            logger.error(f"Background: AI Strategist failed to start: {e}")

        log_event("pagani.api", "system_startup", metadata={
            "timestamp": SERVER_START_TIME.isoformat()
        })
        logger.info("Background Initialization Complete.")

    # Launch background task and return lifespan control immediately
    asyncio.create_task(_background_init())
    
    logger.info("API server yielding to uvicorn (instant startup mode enabled).")
    yield
    logger.info("API server shutting down.")


# ── App ──
app = FastAPI(
    title="Pagani Zonda R – Enterprise Intelligence API",
    description="RAG-powered AI assistant for Pagani Zonda R enterprise data.",
    version="2.0.0",
    lifespan=lifespan,
)

# Rate limiter state
app.state.limiter = limiter


# ── Rate Limit Error Handler ──
@app.exception_handler(RateLimitExceeded)
async def rate_limit_handler(request: Request, exc: RateLimitExceeded):
    return JSONResponse(
        status_code=status.HTTP_429_TOO_MANY_REQUESTS,
        content={
            "detail": "Too many requests. Please slow down.",
            "error_code": "RATE_LIMIT_EXCEEDED",
        },
    )


# ── Global Exception Handler ──
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error(f"Unhandled exception: {exc}", exc_info=True)
    log_event("pagani.api", "api_error", metadata={
        "error": str(exc),
        "path": str(request.url.path),
    })
    return JSONResponse(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        content={
            "detail": "An internal server error occurred.",
            "error_code": "INTERNAL_ERROR",
        },
    )


# ── Register Enterprise Error Handlers ──
register_error_handlers(app)

# ── Security Middleware ──
app.add_middleware(RequestTracingMiddleware)
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(RequestSizeLimitMiddleware)

# ── CORS ──
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", 
        "http://127.0.0.1:3000",
        os.getenv("FRONTEND_URL", ""),
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Request-Id"],
)

# ── Static Files (Serve PDFs via endpoint below, not StaticFiles, to ensure CORS) ──
_pdf_dir = os.path.join(os.path.dirname(__file__), "..", "pagani_intelligence_rich_dataset_25_pdfs")


# ═══════════════════════════════════════════
# Analytics Helper
# ═══════════════════════════════════════════

def _track_analytics(event_type: str, user_id: str | None = None, metadata: dict | None = None):
    """Track a usage analytics event (fire-and-forget)."""
    def _write():
        try:
            from database import get_db_session
            from models import AnalyticsEvent
            with get_db_session() as db:
                db.add(AnalyticsEvent(
                    event_type=event_type,
                    user_id=user_id,
                    metadata_=metadata,
                ))
        except Exception as e:
            logger.warning(f"Analytics tracking failed (non-fatal): {e}")

    # Use a background task for fire-and-forget recording
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            asyncio.create_task(asyncio.to_thread(_write))
        else:
            _write()
    except Exception:
        # Fallback to sync write if no loop is available (e.g. startup)
        _write()


# ═══════════════════════════════════════════
# Chat Persistence Helper
# ═══════════════════════════════════════════

async def _persist_chat(username: str, question: str, response: str):
    """Persist a chat Q&A pair to the database (async-safe)."""
    def _sync_persist():
        try:
            from database import get_db_session
            from models import ChatHistory, User
            with get_db_session() as db:
                user = db.query(User).filter(User.name == username).first()
                if user:
                    db.add(ChatHistory(
                        user_id=user.id,
                        question=question,
                        response=response,
                    ))
        except Exception as e:
            logger.warning(f"Chat persistence failed: {e}")
            
    await asyncio.to_thread(_sync_persist)


# ═══════════════════════════════════════════
# Health Check (Enhanced)
# ═══════════════════════════════════════════

@app.get("/api/health")
async def health_check():
    # Database status
    db_connected = check_db_connection()

    # AI service status
    ai_available = vector_store._initialized

    # Uptime
    uptime_seconds = 0
    if SERVER_START_TIME:
        uptime_seconds = (datetime.now(timezone.utc) - SERVER_START_TIME).total_seconds()

    overall_status = "healthy" if (db_connected and ai_available) else "degraded"

    health_data = {
        "status": overall_status,
        "database": "connected" if db_connected else "disconnected",
        "ai_service": "available" if ai_available else "unavailable",
        "uptime_s": int(uptime_seconds),
        "vector_store_initialized": ai_available,
        "registered_users": len(users_db),
    }

    log_event("pagani.api", "system_health_check", user_id="System", metadata=health_data)

    return {
        **health_data,
        "uptime": f"{uptime_seconds:.0f}s",
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.get("/api/health/detailed")
async def health_check_detailed(current_user: dict = Depends(require_permission("manage_users"))):
    """Secure detailed health check for admins."""
    return await health_check()


# ═══════════════════════════════════════════
# Auth Endpoints
# ═══════════════════════════════════════════

@app.post("/api/register", response_model=dict, status_code=status.HTTP_201_CREATED)
@limiter.limit("10/minute")
async def register(request: Request, user: UserRegister):
    """Register a new user with username, password, and role."""
    logger.info(f"Registration attempt: {user.username} (role: {user.role})")
    result = await register_user(user)
    _track_analytics("user_registered", user_id=user.username, metadata={"role": user.role})
    return {"message": "User registered successfully", **result}


@app.post("/api/login", response_model=TokenResponse)
@limiter.limit("5/minute")
async def login(request: Request, user: UserLogin):
    """Authenticate and receive JWT access + refresh tokens."""
    logger.info(f"Login attempt: {user.username}")
    log_event("pagani.api", "user_login", user_id=user.username)
    result = await authenticate_user(user)
    _track_analytics("login_success", user_id=user.username)
    track_session_start(user.username)
    return result


@app.post("/api/refresh", response_model=TokenResponse)
@limiter.limit("10/minute")
async def refresh(request: Request, body: RefreshRequest):
    """Exchange a valid refresh token for a new access + refresh token pair."""
    logger.info("Token refresh attempt")
    return refresh_access_token(body.refresh_token)


@app.post("/api/logout")
async def logout_endpoint(current_user: dict = Depends(get_current_user)):
    """Log out the current user and track session end."""
    username = current_user["username"]
    logger.info(f"Logout: {username}")
    _track_analytics("logout", user_id=username)
    track_session_end(username)
    return {"message": "Logged out successfully"}


@app.get("/api/me", response_model=UserInfo)
async def get_me(current_user: dict = Depends(get_current_user)):
    """Get current authenticated user info."""
    db_user = users_db.get(current_user["username"])
    return UserInfo(
        username=current_user["username"],
        role=current_user["role"],
        created_at=db_user.get("created_at", "unknown"),
    )


# ═══════════════════════════════════════════
# RAG Chat Endpoints
# ═══════════════════════════════════════════

@app.post("/api/chat", response_model=ChatResponse)
@limiter.limit("20/minute")
async def chat(
    request: Request,
    body: ChatRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    RAG-powered chat endpoint.
    Embeds question → FAISS search (role-filtered) → Gemini generation.
    """
    start_time = time.time()
    username = current_user["username"]
    user_role = current_user["role"]

    logger.info(f"Chat request | user={username} | role={user_role} | q='{body.question[:80]}'")
    log_event("pagani.api", "chat_request", user_id=username, metadata={"question": body.question[:100]})
    _track_analytics("chat_started", user_id=username)
    _track_analytics("query_submitted", user_id=username, metadata={"question_length": len(body.question)})

    try:
        # Step 0: Cache Check
        cache_key = f"chat:{user_role}:{body.question}"
        cached_result = query_cache.get(cache_key)
        if cached_result:
            logger.info(f"Cache HIT for user {username} | query: '{body.question[:50]}'")
            return ChatResponse(**cached_result)

        # Step 1: Agentic Routing (Decide whether to search and reformulate query)
        update_pipeline_node("planner", "running")
        t_plan = time.time()
        history = _get_history(username)
        router_decision = await agentic_router(body.question, history)
        update_pipeline_node("planner", "done", int((time.time() - t_plan) * 1000))
        log_event("pagani.api", "role_routing", user_id=username, metadata=router_decision)
        
        # Step 2: Conditional Semantic Search
        context_docs = []
        if router_decision.get("needs_search", True):
            update_pipeline_node("data_sources", "running")
            update_pipeline_node("tool_execution", "running")
            t_search = time.time()
            search_query = router_decision.get("search_query") or body.question
            logger.info(f"Router decided to search with query: '{search_query[:50]}'")
            context_docs = await asyncio.to_thread(
                vector_store.search,
                query=search_query,
                top_k=5,
                user_role=user_role,
                filters=router_decision.get("metadata_filters")
            )
            search_dur = int((time.time() - t_search) * 1000)
            update_pipeline_node("data_sources", "done", search_dur)
            update_pipeline_node("tool_execution", "done", search_dur)
            logger.info(f"Retrieved {len(context_docs)} documents for user {username}")
        else:
            logger.info(f"Router decided to skip vector search for user {username}")

        # Step 3: Generate response
        update_pipeline_node("router", "running")
        t_route = time.time()
        # Route locally (just simulating router decision for standard chat)
        _ = router_decision.get("decision", "agent_1")
        update_pipeline_node("router", "done", int((time.time() - t_route) * 1000))

        update_pipeline_node("agent_1", "running")
        t_gen = time.time()
        result = await generate_response(
            question=body.question,
            context_docs=context_docs,
            user_role=user_role,
            username=username,
        )
        update_pipeline_node("agent_1", "done", int((time.time() - t_gen) * 1000))

        latency = time.time() - start_time
        logger.info(
            f"Chat response | user={username} | confidence={result['confidence']} | "
            f"sources={len(result['sources'])} | latency={latency:.2f}s"
        )
        log_event("pagani.api", "chat_response", user_id=username, metadata={
            "confidence": result["confidence"],
            "sources": len(result["sources"]),
            "latency_s": round(latency, 2),
        })
        _track_analytics("response_received", user_id=username, metadata={
            "confidence": result["confidence"],
            "latency_s": round(latency, 2),
            "document_ids": result.get("document_ids", []),
            "avg_reranker_score": result.get("avg_reranker_score", 0),
        })

        if not context_docs:
            _track_analytics("failed_query", user_id=username, metadata={"question": body.question[:100]})

        # Persist chat to DB (additive)
        await _persist_chat(username, body.question, result["answer"])

        chat_response = {
            "answer": result["answer"],
            "sources": result["sources"],
            "confidence": result["confidence_score"] / 100.0,
            "user_role": user_role,
        }

        # Cache the result
        query_cache.set(cache_key, chat_response)

        return ChatResponse(**chat_response)

    except RuntimeError as e:
        err_msg = str(e)
        if "QUOTA_EXCEEDED" in err_msg:
            logger.error(f"Quota Exceeded for user {username}: {e}")
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Your AI quota has been exceeded. Please check your Gemini API plan or billing details in Google AI Studio. (Error 429: Quota Exceeded)",
            )
        
        if "INVALID_API_KEY" in err_msg:
            logger.error(f"Invalid API Key for user {username}: {e}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Your Gemini API key is expired or invalid. Please update GEMINI_API_KEY in your backend/.env file and restart the server.",
            )
        
        logger.error(f"RAG pipeline error for user {username}: {e}")
        log_event("pagani.api", "api_error", user_id=username, metadata={"error": err_msg})
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="The AI service is temporarily unavailable. Please try again.",
        )
    except Exception as e:
        logger.error(f"Unexpected chat error for user {username}: {e}", exc_info=True)
        log_event("pagani.api", "api_error", user_id=username, metadata={"error": str(e)})
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred processing your request.",
        )


@app.post("/api/chat/debug")
@limiter.limit("20/minute")
async def chat_debug(
    request: Request,
    body: ChatRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Debug-enhanced RAG chat endpoint.
    Returns the full pipeline trace alongside the normal response.
    Existing /api/chat and /api/chat/stream are NOT modified.
    """
    import time as _time
    t_start = _time.time()
    username = current_user["username"]
    user_role = current_user["role"]

    logger.info(f"Debug chat request | user={username} | role={user_role} | q='{body.question[:80]}'")

    try:
        # Step 1: Agentic Routing
        history = _get_history(username)
        router_decision = await agentic_router(body.question, history)

        # Step 2: Search with debug info
        context_docs = []
        debug_info = {
            "pipeline_steps": [
                {"step": "query_received", "label": "Query Received", "timestamp_ms": 0}
            ],
            "search_results": [],
            "retrieved_chunks": [],
            "timing": {},
            "router_decision": router_decision,
        }

        if router_decision.get("needs_search", True):
            search_query = router_decision.get("search_query") or body.question

            # Use the debug-enhanced search
            search_result = await asyncio.to_thread(
                vector_store.search_with_debug,
                query=search_query,
                top_k=5,
                user_role=user_role,
                filters=router_decision.get("metadata_filters")
            )
            context_docs = search_result["results"]
            debug_info = search_result["debug"]
            debug_info["router_decision"] = router_decision

        # Step 3: Generate response
        t_gen = _time.time()
        result = await generate_response(
            question=body.question,
            context_docs=context_docs,
            user_role=user_role,
            username=username,
        )
        gen_ms = int((_time.time() - t_gen) * 1000)
        total_ms = int((_time.time() - t_start) * 1000)

        debug_info["timing"]["generation_ms"] = gen_ms
        debug_info["timing"]["total_ms"] = total_ms
        debug_info["pipeline_steps"].append({
            "step": "llm_generated",
            "label": "LLM Response Generated",
            "timestamp_ms": total_ms,
        })

        # Persist chat (additive, same as normal endpoint)
        await _persist_chat(username, body.question, result["answer"])

        return {
            "answer": result["answer"],
            "sources": result["sources"],
            "confidence": result["confidence_score"] / 100.0,
            "user_role": user_role,
            "debug": debug_info,
        }

    except RuntimeError as e:
        logger.error(f"Debug RAG pipeline error for user {username}: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="The AI service is temporarily unavailable. Please try again.",
        )
    except Exception as e:
        logger.error(f"Unexpected debug chat error for user {username}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred processing your request.",
        )


@app.post("/api/chat/stream")
@limiter.limit("20/minute")
async def chat_stream(
    request: Request,
    body: ChatRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Streaming RAG chat endpoint.
    Returns Server-Sent Events with token-by-token response.
    """
    username = current_user["username"]
    user_role = current_user["role"]

    logger.info(f"Stream chat request | user={username} | role={user_role} | q='{body.question[:80]}'")
    log_event("pagani.api", "chat_request", user_id=username, metadata={
        "question": body.question[:100],
        "streaming": True,
    })
    _track_analytics("chat_started", user_id=username, metadata={"streaming": True})

    try:
        # Step 0: Cache Check (for simplicity, only full completions are cached)
        cache_key = f"chat_stream:{user_role}:{body.question}"
        cached_result = query_cache.get(cache_key)
        
        if cached_result:
            logger.info(f"Stream Cache HIT for user {username} | query: '{body.question[:50]}'")
            async def cached_generator():
                yield f"data: {cached_result['answer']}\n\n"
                yield "data: [DONE]\n\n"
            return StreamingResponse(
                cached_generator(),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
            )

        history = _get_history(username)
        router_decision = await agentic_router(body.question, history)
        
        context_docs = []
        if router_decision.get("needs_search", True):
            search_query = router_decision.get("search_query") or body.question
            logger.info(f"Router decided to search with query: '{search_query[:50]}'")
            context_docs = await asyncio.to_thread(
                vector_store.search,
                query=search_query,
                top_k=5,
                user_role=user_role,
                filters=router_decision.get("metadata_filters")
            )
        else:
            logger.info(f"Router decided to skip vector search for user {username}")

        collected_chunks: list[str] = []

        async def event_generator():
            async for chunk in generate_response_stream(
                question=body.question,
                context_docs=context_docs,
                user_role=user_role,
                username=username,
            ):
                collected_chunks.append(chunk)
                yield f"data: {chunk}\n\n"
            yield "data: [DONE]\n\n"

            # Persist after streaming completes
            full_response = "".join(collected_chunks)
            await _persist_chat(username, body.question, full_response)
            
            # Cache the result
            query_cache.set(cache_key, {"answer": full_response})
            
            _track_analytics("response_received", user_id=username, metadata={
                "streaming": True,
                "ttft_ms": int(ttft * 1000) if 'ttft' in locals() else None,
                "document_ids": list({doc.get("doc_id", doc.get("source", "unknown")) for doc in context_docs}),
            })
            if not context_docs:
                _track_analytics("failed_query", user_id=username, metadata={"question": body.question[:100]})
            log_event("pagani.api", "chat_response", user_id=username, metadata={"streaming": True})

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )

    except Exception as e:
        logger.error(f"Streaming RAG error for user {username}: {e}")
        log_event("pagani.api", "api_error", user_id=username, metadata={"error": str(e)})
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="The AI service is temporarily unavailable.",
        )


@app.get("/api/chat/quick-summary", summary="Get a quick summary for a model")
@limiter.limit("30/minute")
async def chat_quick_summary(
    request: Request,
    model: str = QueryParam(..., description="Vehicle model name"),
    current_user: dict = Depends(get_current_user),
):
    """Return a brief conceptual summary for the requested Pagani model."""
    summaries = {
        "zonda_r": "The Zonda R is a track-only hypercar developed purely for performance, featuring a carbon-titanium chassis and a 750hp AMG V12.",
        "huayra_bc": "The Huayra BC honors Benny Caiola, introducing lightweight aerodynamics and an upgraded twin-turbo V12 packing 789hp.",
        "utopia": "Utopia, the third act of Pagani, embraces mechanical purity with its manual transmission, twin-turbo V12, and classic analog aesthetics."
    }
    
    model_key = model.lower().replace(" ", "_")
    for key, text in summaries.items():
        if key in model_key or model_key in key:
            return {"model": model, "summary": text}
            
    return {"model": model, "summary": f"The {model} showcases Pagani's signature blend of art and science, combining bespoke engineering with exquisite craftsmanship."}


@app.post("/api/chat/suggestions", summary="Get chat suggestions")
@limiter.limit("30/minute")
async def chat_suggestions(
    request: Request,
    current_user: dict = Depends(get_current_user),
):
    """Provide intelligent chat suggestions based on user role."""
    role = current_user.get("role", "viewer")
    
    if role == "admin":
        suggestions = [
            "What is the total revenue from Zonda R sales?",
            "What is the current market valuation of the Zonda R?",
            "Show me the production timeline for all 15 units."
        ]
    elif role == "engineer":
        suggestions = [
            "What is the torsional rigidity of the Zonda R monocoque?",
            "Detail the Öhlins damper specifications.",
            "What is the peak downforce at 300 km/h?"
        ]
    else:
        suggestions = [
            "What is the top speed of the Zonda R?",
            "Tell me about the carbon-titanium chassis.",
            "How much horsepower does the AMG V12 produce?"
        ]
        
    return {"suggestions": suggestions}


# ═══════════════════════════════════════════
# V1 Enterprise API Router
# ═══════════════════════════════════════════

v1_router = APIRouter(prefix="/api/v1", tags=["Enterprise V1"])


# ── Pydantic Models for V1 ──
class RoleChangeRequest(BaseModel):
    new_role: str = Field(..., description="New role to assign")

    @field_validator("new_role")
    @classmethod
    def sanitize_new_role(cls, v: str) -> str:
        return bleach.clean(v, tags=[], strip=True)

class DocumentMetadataUpdate(BaseModel):
    title: Optional[str] = None
    tags: Optional[list[str]] = None

    @field_validator("title")
    @classmethod
    def sanitize_title(cls, v: Optional[str]) -> Optional[str]:
        if v:
            return bleach.clean(v, tags=[], strip=True)
        return v

    @field_validator("tags")
    @classmethod
    def sanitize_tags(cls, v: Optional[list[str]]) -> Optional[list[str]]:
        if v:
            return [bleach.clean(tag, tags=[], strip=True) for tag in v]
        return v


# ───────────────────────────
# RBAC Admin Endpoints
# ───────────────────────────

@v1_router.get("/admin/users", summary="List all users")
async def v1_list_users(
    current_user: dict = Depends(require_permission("manage_users")),
):
    """List all registered users (admin/super_admin only)."""
    return {
        "users": [
            {"username": u, "role": d["role"], "created_at": d.get("created_at", "unknown")}
            for u, d in users_db.items()
        ],
        "total": len(users_db),
    }


@v1_router.delete("/admin/users/{username}", summary="Delete user account")
async def v1_delete_user(
    username: str,
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Delete a user account and logically remove from the database."""
    if username not in users_db:
        raise HTTPException(status_code=404, detail="User not found")
        
    if username == current_user["username"]:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")

    # Remove from DB
    try:
        from database import get_db_session
        from models import User
        with get_db_session() as db:
            user = db.query(User).filter(User.name == username).first()
            if user:
                db.delete(user)
    except Exception as e:
        logger.error(f"Failed to delete user from database: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete user")

    # Remove from active memory
    del users_db[username]
    
    # Audit log
    audit.log("user_deleted", current_user["username"], metadata={"deleted_user": username})
    return {"message": f"User {username} deleted successfully"}


@v1_router.put("/admin/users/{username}/role", summary="Change user role")
async def v1_change_user_role(
    username: str,
    body: RoleChangeRequest,
    current_user: dict = Depends(require_permission("manage_roles")),
):
    """Change a user's role (super_admin only)."""
    if username not in users_db:
        raise HTTPException(status_code=404, detail="User not found")
    if body.new_role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {', '.join(VALID_ROLES)}")

    old_role = users_db[username]["role"]
    users_db[username]["role"] = body.new_role

    # Log role change
    audit.log_role_change(
        changed_by=current_user["username"],
        target_user=username,
        old_role=old_role,
        new_role=body.new_role,
    )

    # Persist to DB
    try:
        from database import get_db_session
        from models import User, RoleAuditLog
        with get_db_session() as db:
            user = db.query(User).filter(User.name == username).first()
            if user:
                user.role = body.new_role
            db.add(RoleAuditLog(
                changed_by=current_user["username"],
                target_user=username,
                old_role=old_role,
                new_role=body.new_role,
            ))
    except Exception as e:
        logger.warning(f"Role change DB persistence failed: {e}")

    return {
        "message": f"Role updated: {username} ({old_role} -> {body.new_role})",
        "username": username,
        "old_role": old_role,
        "new_role": body.new_role,
    }


@v1_router.get("/admin/roles/audit", summary="Role change audit log")
async def v1_role_audit_log(
    limit: int = QueryParam(default=50, le=500),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """View role change audit trail."""
    try:
        from database import get_db_session
        from models import RoleAuditLog
        with get_db_session() as db:
            logs = db.query(RoleAuditLog).order_by(RoleAuditLog.timestamp.desc()).limit(limit).all()
            formatted_logs = [
                {
                    "id": log.id,
                    "changed_by": log.changed_by,
                    "target_user": log.target_user,
                    "old_role": log.old_role,
                    "new_role": log.new_role,
                    "timestamp": log.timestamp.isoformat() if log.timestamp else None,
                }
                for log in logs
            ]
            return {"audit_logs": formatted_logs, "total": len(formatted_logs)}
    except Exception as e:
        logger.error(f"Failed to retrieve role audit logs: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve role audit logs")


@v1_router.get("/admin/permissions", summary="View permission matrix")
async def v1_permissions(
    current_user: dict = Depends(get_current_user),
):
    """View the RBAC permission matrix."""
    return {
        "permissions": ROLE_PERMISSIONS,
        "valid_roles": list(VALID_ROLES),
        "your_role": current_user["role"],
        "your_permissions": ROLE_PERMISSIONS.get(current_user["role"], []),
    }


# ───────────────────────────
# Analytics Endpoints
# ───────────────────────────

@v1_router.get("/analytics/engagement", summary="User engagement metrics")
async def v1_analytics_engagement(
    days: int = QueryParam(default=30, le=365),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Get user engagement metrics for the specified period."""
    return get_user_engagement_metrics(days=days)


@v1_router.get("/analytics/queries", summary="Query success/failure rates")
async def v1_analytics_queries(
    days: int = QueryParam(default=30, le=365),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Get query success/failure rate statistics."""
    return get_query_success_rates(days=days)


@v1_router.get("/analytics/ai-performance", summary="AI performance metrics")
async def v1_analytics_ai(
    days: int = QueryParam(default=30, le=365),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Get AI performance metrics (confidence, latency)."""
    return get_ai_performance_metrics(days=days)


@v1_router.get("/analytics/system-health", summary="System health metrics")
async def v1_system_health(
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Get system health metrics (CPU, memory, uptime)."""
    return get_system_health()


@v1_router.get("/analytics/summary", summary="Aggregated analytics summary")
async def v1_analytics_summary(
    days: int = QueryParam(default=7, le=30),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Get aggregated metrics for the admin dashboard."""
    return get_analytics_summary(days=days)


@v1_router.get("/analytics/export", summary="Export analytics as CSV")
async def v1_analytics_export(
    days: int = QueryParam(default=30, le=365),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Export analytics events as CSV."""
    csv_data = export_analytics_csv(days=days)
    return PlainTextResponse(
        content=csv_data,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=analytics_{days}d.csv"},
    )


# ───────────────────────────
# Audit Endpoints
# ───────────────────────────

@v1_router.get("/audit/logs", summary="View audit logs")
async def v1_audit_logs(
    action: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = QueryParam(default=100, le=1000),
    offset: int = QueryParam(default=0, ge=0),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Admin view of all audit/system logs."""
    logs = get_audit_logs(action=action, user_id=user_id, limit=limit, offset=offset)
    return {"logs": logs, "total": len(logs)}


@v1_router.get("/audit/login-attempts", summary="Login attempt history")
async def v1_login_attempts(
    limit: int = QueryParam(default=50, le=500),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """View recent login attempts."""
    attempts = get_login_attempts(limit=limit)
    return {"attempts": attempts, "total": len(attempts)}


# ───────────────────────────
# Recent Chat History (RAG Query Audit)
# ───────────────────────────

@v1_router.get("/admin/recent-chats", summary="Recent RAG chat queries")
async def v1_recent_chats(
    limit: int = QueryParam(default=10, le=50),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Fetch the most recent chat queries with user info from the database."""
    try:
        from database import get_db_session
        from models import ChatHistory, User
        with get_db_session() as db:
            rows = (
                db.query(ChatHistory, User)
                .join(User, ChatHistory.user_id == User.id)
                .order_by(ChatHistory.timestamp.desc())
                .limit(limit)
                .all()
            )
            chats = []
            for chat, user in rows:
                ts = chat.timestamp
                time_str = ts.strftime("%H:%M") if ts else ""
                chats.append({
                    "id": chat.id,
                    "user": user.name,
                    "role": user.role,
                    "question": chat.question,
                    "time": time_str,
                    "timestamp": ts.isoformat() if ts else None,
                })
            return {"chats": chats, "total": len(chats)}
    except Exception as e:
        logger.error(f"Failed to fetch recent chats: {e}")
        return {"chats": [], "total": 0}


# ───────────────────────────
# Document Management Endpoints
# ───────────────────────────

@v1_router.get("/static/pdfs/{filename}", summary="Serve static PDF files")
async def serve_pdf(filename: str):
    """Serve a static PDF document from the dataset directory."""
    file_path = os.path.join(_pdf_dir, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="PDF not found")
    return FileResponse(file_path, media_type="application/pdf")

@v1_router.get("/models", summary="List Pagani models")
async def get_models(current_user: dict = Depends(get_current_user)):
    """Return the list of Pagani models for the showcase and comparison tools."""
    models = [
        { 
            "name": "Zonda R", "hp": 740, "weight": 1070, "badge": "Featured", "badgeColor": "bg-pagani-gold/15 text-pagani-gold border-pagani-gold/25", "featured": True, "imageUrl": "/images/models/zonda_r.png", 
            "summary": "The Zonda R is a track-only hypercar designed to be the ultimate expression of Pagani's engineering capabilities. Free from road and racing regulations, it features a bespoke carbon-titanium chassis and a blistering aerodynamic package. The naturally aspirated 6.0L V12 derived from the Mercedes-Benz CLK GTR produces a deafening roar, delivering a pure, visceral driving experience that set a Nürburgring Nordschleife record of 6:47.48.",
            "engine": "Mercedes-Benz AMG 6.0L V12",
            "topSpeed": "350+ km/h (218 mph)",
            "acceleration": "2.7s (0-100 km/h)",
            "productionUnits": 15,
            "price": "€1.4 Million"
        },
        { 
            "name": "Huayra BC", "hp": 789, "weight": 1218, "badge": "New doc", "badgeColor": "bg-green-500/15 text-green-400 border-green-500/25", "featured": False, "imageUrl": "/images/models/huayra_bc.png", 
            "summary": "Named after Horacio Pagani’s first customer and close friend, Benny Caiola, the Huayra BC is a more aggressive, track-focused evolution of the standard Huayra. It introduces a revolutionary new carbon-triax composite that is 20% stronger and 50% lighter than regular carbon fiber. Combined with a significantly uprated twin-turbo V12 and an entirely new active aerodynamic system, it generates immense downforce while shedding over 130 kg from the base model.",
            "engine": "Mercedes-AMG 6.0L Twin-Turbo V12",
            "topSpeed": "370 km/h (230 mph)",
            "acceleration": "2.8s (0-100 km/h)",
            "productionUnits": 20,
            "price": "€2.3 Million"
        },
        { 
            "name": "Utopia", "hp": 864, "weight": 1280, "badge": "New", "badgeColor": "bg-sky-500/15 text-sky-400 border-sky-500/25", "featured": False, "imageUrl": "/images/models/utopia.png", 
            "summary": "The Utopia represents the third major chapter in Pagani's history. Eschewing the modern trend of heavy hybrid systems and dual-clutch transmissions, the Utopia focuses on purity, offering a breathtaking 864 hp AMG V12 paired with a 7-speed gated manual transmission. Its design blends retro-futuristic elegance with active aerodynamics subtly integrated into the bodywork, delivering a timeless aesthetic and an engaging, analog driving experience.",
            "engine": "Mercedes-AMG 6.0L Twin-Turbo V12",
            "topSpeed": "350 km/h (217 mph) electronically limited",
            "acceleration": "2.9s (0-100 km/h)",
            "productionUnits": 99,
            "price": "€2.17 Million"
        },
        { 
            "name": "Zonda F", "hp": 602, "weight": 1230, "badge": "Classic", "badgeColor": "bg-amber-500/15 text-amber-400 border-amber-500/25", "featured": False, "imageUrl": "/images/models/zonda_f.png", 
            "summary": "Dedicated to five-time Formula One World Champion Juan Manuel Fangio, the Zonda F is widely considered the definitive road-going Zonda. It refined the original Zonda formula with improved aerodynamics, a redesigned front fascia, and a carbon-ceramic braking system. The glorious naturally aspirated V12 provides instant throttle response, making it one of the most highly sought-after collector cars of the modern era.",
            "engine": "Mercedes-Benz AMG 7.3L V12",
            "topSpeed": "345 km/h (214 mph)",
            "acceleration": "3.6s (0-100 km/h)",
            "productionUnits": 25,
            "price": "€1.0 Million (Original)"
        },
        { 
            "name": "Cinque", "hp": 678, "weight": 1210, "badge": "Rare", "badgeColor": "bg-purple-500/15 text-purple-400 border-purple-500/25", "featured": False, "imageUrl": "/images/models/cinque.png", 
            "summary": "Originally built at the request of the Pagani dealer in Hong Kong, the Zonda Cinque was the first road-legal car to feature Pagani's revolutionary carbon-titanium weave. It borrows heavily from the track-only Zonda R, incorporating its aggressive roof scoop, revised aerodynamics, and magnesium wheels. As the name suggests, only five coupes were ever produced, making it a true unicorn in the automotive world.",
            "engine": "Mercedes-Benz AMG 7.3L V12",
            "topSpeed": "350 km/h (217 mph)",
            "acceleration": "3.4s (0-100 km/h)",
            "productionUnits": 5,
            "price": "€1.3 Million (Original)"
        },
    ]
    return {"models": models}

@v1_router.get("/documents/count", summary="Get document count for role")
async def v1_count_documents(
    role: str = QueryParam("viewer", description="Role to filter by"),
    current_user: dict = Depends(get_current_user),
):
    """Count documents accessible to a given role."""
    count = 0
    docs = getattr(vector_store, "documents", [])
    for d in docs:
        if role in d.get("role_access", ["viewer", "seller", "admin"]):
            count += 1
    return {"count": count, "role": role}


@v1_router.get("/documents", summary="List all documents")
async def v1_list_documents(
    role: Optional[str] = QueryParam(None, description="Filter by role"),
    limit: int = QueryParam(50, description="Max documents to return"),
    sort: str = QueryParam("recent", description="Sort order"),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """List all documents in the vector store with optional filtering."""
    try:
        # Prefer database for managed documents
        db_docs = db.query(Document).all()
        
        # Merge with vector store documents if any (e.g. static ones)
        vs_docs = getattr(vector_store, "documents", [])
        
        # Use a dict to merge by filename/doc_id
        merged = {}
        
        # 1. Add DB docs
        for d in db_docs:
            merged[d.filename] = {
                "id": d.id,
                "filename": d.filename,
                "type": d.file_type,
                "file_size": d.file_size,
                "version": d.version,
                "uploaded_by": d.uploaded_by,
                "created_at": d.created_at.isoformat() if d.created_at else None,
                "updated_at": d.updated_at.isoformat() if d.updated_at else None,
            }
            
        # 2. Add VS docs if not already present
        for i, d in enumerate(vs_docs):
            fname = d.get("source", f"vs_{i}")
            if fname not in merged:
                merged[fname] = {
                    "id": d.get("doc_id", str(i)),
                    "filename": fname,
                    "type": "PDF" if d.get("is_pdf") else "SPEC",
                    "file_size": len(d.get("content", "")),
                    "version": "1",
                    "uploaded_by": "System",
                    "created_at": d.get("created_at", "2026-04-01T10:00:00Z"),
                }

        docs_list = list(merged.values())
        
        # Filter by role (if needed - currently we don't have per-doc RBAC in DB yet)
        # For now, just return all if admin/engineer
        
        # Sort
        if sort == "recent":
            docs_list.sort(key=lambda x: x.get("created_at", "") or "", reverse=True)
            
        # Limit
        total = len(docs_list)
        docs_list = docs_list[:limit]
        
        return {"documents": docs_list, "total": total}
    except Exception as e:
        logger.warning(f"Failed to list documents: {e}")
        return {"documents": [], "total": 0}


@v1_router.get("/documents/{doc_id}", summary="Get document by ID")
async def v1_get_document(
    doc_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Retrieve a single document by its ID."""
    try:
        doc = vector_store.get_document(doc_id) if hasattr(vector_store, "get_document") else None
        if doc is None:
            raise HTTPException(status_code=404, detail="Document not found")
        return doc
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"Failed to get document {doc_id}: {e}")
        raise HTTPException(status_code=404, detail="Document not found")


@v1_router.delete("/documents/{doc_id}", summary="Delete a document")
async def v1_delete_document(
    doc_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Delete a document from both the database and the vector store."""
    try:
        # 1. Delete from Database
        doc = db.query(Document).filter(Document.id == doc_id).first()
        db_deleted = False
        if doc:
            # Also delete physical file if it exists
            if doc.file_path and os.path.exists(doc.file_path):
                try:
                    os.remove(doc.file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete physical file {doc.file_path}: {e}")
            
            # Delete associated versions
            db.query(DocumentVersion).filter(DocumentVersion.document_id == doc_id).delete()
            
            db.delete(doc)
            db.commit()
            db_deleted = True
            logger.info(f"Deleted document {doc_id} from database")

        # 2. Delete from Vector Store
        vs_deleted = vector_store.delete_document(doc_id) if hasattr(vector_store, "delete_document") else False
        
        # Fallback: if doc_id was a filename, try that too
        if not vs_deleted and doc:
            vs_deleted = vector_store.delete_document(doc.filename) if hasattr(vector_store, "delete_document") else False

        if not db_deleted and not vs_deleted:
            raise HTTPException(status_code=404, detail="Document not found in database or vector store")

        audit.log("document_deleted", current_user["username"], metadata={"doc_id": doc_id})
        return {"message": f"Document {doc_id} deleted successfully", "id": doc_id}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to delete document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


@v1_router.post("/documents/upload", summary="Upload a document")
async def v1_upload_document(
    request: Request,
    db: Session = Depends(get_db),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Upload and ingest a document into the vector store with versioning."""
    form = await request.form()
    file = form.get("file")
    if file is None:
        raise HTTPException(status_code=422, detail="No file provided")

    ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".json", ".docx"}
    filename = getattr(file, "filename", "") or ""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    try:
        content = await file.read()
        content_hash = hashlib.sha256(content).hexdigest()
        
        # Check if document already exists by filename
        doc = db.query(Document).filter(Document.filename == filename).first()
        if not doc:
            doc = Document(
                filename=filename,
                file_type=ext.replace(".", "").upper(),
                file_size=str(len(content)),
                uploaded_by=current_user["username"],
                version="1"
            )
            db.add(doc)
            db.flush() # Get doc.id
            version_num = 1
        else:
            # Increment version
            version_num = int(doc.version) + 1
            doc.version = str(version_num)
            doc.file_size = str(len(content))
            doc.updated_at = datetime.now(timezone.utc)
        
        # Save file to disk for version tracking
        # Ensure directory exists in workspace
        upload_dir = os.path.join(os.path.dirname(__file__), "data", "uploads", "versions")
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, f"{doc.id}_v{version_num}{ext}")
        with open(file_path, "wb") as f:
            f.write(content)
        
        doc.file_path = file_path
        
        # Create version record
        doc_version = DocumentVersion(
            document_id=doc.id,
            version_number=version_num,
            content_hash=content_hash,
            file_path=file_path,
            created_by=current_user["username"]
        )
        db.add(doc_version)
        db.commit()

        # Delegate to vector store ingestion if available
        if hasattr(vector_store, "ingest_document"):
            vector_store.ingest_document(filename, content)

        audit.log("document_uploaded", current_user["username"], metadata={
            "filename": filename, "size_bytes": len(content), "version": version_num
        })
        return {"message": "Document uploaded successfully", "id": doc.id, "version": version_num}
    except Exception as e:
        db.rollback()
        logger.error(f"Document upload failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to process uploaded document")


@v1_router.get("/documents/{doc_id}/versions", summary="List document versions")
async def v1_list_document_versions(
    doc_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """List all versions of a specific document."""
    versions = db.query(DocumentVersion).filter(DocumentVersion.document_id == doc_id).order_by(DocumentVersion.version_number.desc()).all()
    # Format for JSON
    formatted = []
    for v in versions:
        formatted.append({
            "id": v.id,
            "version_number": v.version_number,
            "created_at": v.created_at.isoformat() if v.created_at else None,
            "created_by": v.created_by,
            "hash": v.content_hash[:8] if v.content_hash else "N/A"
        })
    return {"versions": formatted}


@v1_router.post("/documents/{doc_id}/restore/{version_number}", summary="Restore document version")
async def v1_restore_document_version(
    doc_id: str,
    version_number: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Restore a document to a specific version."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    version = db.query(DocumentVersion).filter(
        DocumentVersion.document_id == doc_id,
        DocumentVersion.version_number == version_number
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    # Update main document to point to this version's file/metadata
    doc.version = str(version_number)
    doc.file_path = version.file_path
    db.commit()
    
    # Re-ingest the version content into vector store to reflect changes in RAG
    if os.path.exists(version.file_path) and hasattr(vector_store, "ingest_document"):
        with open(version.file_path, "rb") as f:
            content = f.read()
            vector_store.ingest_document(doc.filename, content)
    
    audit.log("document_restored", current_user["username"], metadata={
        "doc_id": doc_id, "version": version_number
    })
    return {"message": f"Document restored to version {version_number}", "id": doc_id}


@v1_router.get("/documents/{doc_id}/content", summary="Get active document content")
async def v1_get_document_content(
    doc_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """Retrieve the text content of the active version of a document (DB or VS)."""
    logger.info(f"Retrieving content for doc_id: {doc_id}")
    
    # 1. Try Database lookup first
    doc = db.query(Document).filter(Document.id == doc_id).first()
    
    if doc:
        if not doc.file_path:
            logger.warning(f"Document {doc_id} has no file_path recorded")
            raise HTTPException(status_code=404, detail="No file path associated with this document")
        
        if not os.path.exists(doc.file_path):
            logger.warning(f"Physical file missing for doc {doc_id}: {doc.file_path}. Attempting VectorStore fallback...")
            
            # Try to recover content from VectorStore (merging chunks if necessary)
            vs_doc = vector_store.get_document(doc_id) or vector_store.get_document(doc.filename)
            
            if vs_doc and vs_doc.get("content"):
                logger.info(f"Successfully recovered content from VectorStore for {doc.filename}")
                return {
                    "content": vs_doc["content"], 
                    "filename": doc.filename,
                    "note": "Recovered from vector store (original file missing)"
                }
                
            # If fallback fails, raise descriptive error
            logger.error(f"Recovery failed for doc {doc_id}")
            is_temp = "Temp" in doc.file_path or "tmp" in doc.file_path.lower()
            detail = f"Physical file missing: {os.path.basename(doc.file_path)}"
            if is_temp:
                detail = f"Document '{doc.filename}' was stored in temporary storage and has been cleared by the system. Please re-upload it."
            raise HTTPException(status_code=404, detail=detail)
        
        try:
            # Simple text extraction based on extension
            ext = doc.filename.split('.')[-1].lower()
            if ext == 'txt':
                with open(doc.file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif ext == 'pdf':
                import fitz  # pymupdf
                doc_pdf = fitz.open(doc.file_path)
                content = "".join([page.get_text() for page in doc_pdf])
                doc_pdf.close()
            elif ext in ['docx', 'doc']:
                import docx
                doc_word = docx.Document(doc.file_path)
                content = "\n".join([p.text for p in doc_word.paragraphs])
            else:
                # Fallback to reading as bytes and decode
                with open(doc.file_path, "rb") as f:
                    content = f.read().decode('utf-8', errors='ignore')
            
            return {"content": content, "filename": doc.filename}
        except Exception as e:
            logger.error(f"Failed to read database document content for {doc_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to read document content: {str(e)}")

    # 2. Fallback to Vector Store for hardcoded/system docs
    vs_doc = vector_store.get_document(doc_id)
    if vs_doc:
        return {
            "content": vs_doc.get("content", ""),
            "filename": vs_doc.get("filename", "System Document")
        }
    
    logger.warning(f"Document not found in DB or VS: {doc_id}")
    raise HTTPException(status_code=404, detail="Document not found")




# ───────────────────────────
# Cache Stats Endpoint
# ───────────────────────────

@v1_router.get("/cache/stats", summary="Cache statistics")
async def v1_cache_stats(
    current_user: dict = Depends(require_permission("manage_users")),
):
    """View query cache statistics."""
    return {
        "query_cache": query_cache.stats,
    }


# ───────────────────────────
# WebSocket Endpoints
# ───────────────────────────

@v1_router.websocket("/ws/notifications")
async def ws_notifications(websocket: WebSocket):
    """WebSocket for real-time notifications."""
    await ws_manager.connect(websocket, "notifications")
    try:
        while True:
            # Keep connection alive, optionally receive client messages
            data = await websocket.receive_text()
            # Echo or handle client messages
            await ws_manager.send_personal(websocket, {"type": "ack", "message": data})
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket, "notifications")


@v1_router.websocket("/ws/logs")
async def ws_logs(websocket: WebSocket):
    """WebSocket for real-time log streaming (admin only)."""
    await ws_manager.connect(websocket, "logs")
    try:
        while True:
            data = await websocket.receive_text()
            await ws_manager.send_personal(websocket, {"type": "ack", "message": data})
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket, "logs")




# ───────────────────────────
# Evaluation Endpoints
# ───────────────────────────

class EvalRequest(BaseModel):
    query: str
    response: str
    reference: Optional[str] = None
    retrieved_ids: Optional[list[str]] = None
    relevant_ids: Optional[list[str]] = None
    input_tokens: int = 0
    output_tokens: int = 0
    latency_seconds: float = 0.0

    @field_validator("query", "response", "reference")
    @classmethod
    def sanitize_strings(cls, v: Optional[str]) -> Optional[str]:
        if v:
            return bleach.clean(v, tags=[], strip=True)
        return v

@v1_router.post("/evaluate", summary="Evaluate a RAG response")
async def v1_evaluate(
    body: EvalRequest,
    current_user: dict = Depends(require_permission("execute")),
):
    """Run LLMJudge scoring + IR metrics + cost tracking and persist results."""
    evaluator = Evaluator()
    result = evaluator.evaluate(
        query=body.query,
        response=body.response,
        reference=body.reference,
        retrieved_ids=body.retrieved_ids,
        relevant_ids=body.relevant_ids,
        input_tokens=body.input_tokens,
        output_tokens=body.output_tokens,
        latency_seconds=body.latency_seconds,
    )
    return result


@v1_router.post("/evaluate/ir-metrics", summary="Compute IR metrics only")
async def v1_ir_metrics(
    retrieved_ids: list[str],
    relevant_ids: list[str],
    current_user: dict = Depends(get_current_user),
):
    """Compute Precision, Recall, F1 without LLM judge."""
    return IRMetrics.compute(retrieved_ids, relevant_ids)


# ───────────────────────────
# Admin – Review Queue / Feedback / Golden Answers
# ───────────────────────────

@v1_router.get("/admin/review-queue", summary="View flagged queries")
async def v1_review_queue(
    _key: str = Depends(verify_admin_key),
):
    """View all Gatekeeper-flagged queries awaiting review (X-Admin-Key)."""
    items = []
    try:
        from database import get_db_session
        from models import ReviewQueue as RQ
        with get_db_session() as db:
            rows = db.query(RQ).filter(RQ.status == "pending_review").order_by(RQ.created_at.desc()).all()
            items = [{"id": r.id, "username": r.username, "question": r.question, "reason": r.reason, "status": r.status, "created_at": r.created_at.isoformat() if r.created_at else None} for r in rows]
    except Exception:
        items = [{"id": k, **v} for k, v in auth_review_queue.items()]
    return {"review_queue": items, "total": len(items)}


class ReviewActionRequest(BaseModel):
    action: str = Field(..., pattern="^(approve|reject|edit)$")
    edited_response: Optional[str] = None

    @field_validator("action", "edited_response")
    @classmethod
    def sanitize_strings(cls, v: Optional[str]) -> Optional[str]:
        if v:
            return bleach.clean(v, tags=[], strip=True)
        return v


@v1_router.patch("/admin/review/{item_id}", summary="Approve/reject/edit a flagged query")
async def v1_review_action(
    item_id: str,
    body: ReviewActionRequest,
    _key: str = Depends(verify_admin_key),
):
    """Resolve a review queue item with approve/reject/edit action (X-Admin-Key)."""
    try:
        from database import get_db_session
        from models import ReviewQueue as RQ
        with get_db_session() as db:
            row = db.query(RQ).filter(RQ.id == item_id).first()
            if not row:
                raise HTTPException(status_code=404, detail="Review item not found")
            if body.action == "approve":
                row.status = "approved"
            elif body.action == "reject":
                row.status = "rejected"
            elif body.action == "edit":
                row.status = "approved"
                if body.edited_response:
                    row.final_response = body.edited_response
            row.resolved_at = datetime.now(timezone.utc)
            db.commit()
            return {"message": f"Review item {body.action}d", "id": item_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"DB resolve failed: {e}")
    if item_id in auth_review_queue:
        auth_review_queue[item_id]["status"] = body.action + "d"
        return {"message": "Resolved (in-memory)", "id": item_id}
    raise HTTPException(status_code=404, detail="Review item not found")


@v1_router.get("/admin/audit-log", summary="View audit logs")
async def v1_admin_audit_log(
    action: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = QueryParam(default=100, le=500),
    offset: int = QueryParam(default=0, ge=0),
    _key: str = Depends(verify_admin_key),
):
    """View all audit log entries (X-Admin-Key)."""
    logs = get_audit_logs(action=action, user_id=user_id, limit=limit, offset=offset)
    return {"logs": logs, "total": len(logs)}


@v1_router.get("/admin/stress/stream", summary="Run stress tests via SSE")
async def v1_stress_test_stream(
    test_type: str = QueryParam(default="all"),
    _key: str = Depends(verify_admin_key)
):
    """Run adversarial safety tests and stream results (X-Admin-Key)."""
    tester = StressTester()
    if test_type == "bias":
        generator = tester.run_bias_test_stream()
    elif test_type == "evasion":
        generator = tester.run_evasion_test_stream()
    elif test_type == "injection":
        generator = tester.run_injection_test_stream()
    else:
        generator = tester.run_all_stream()
        
    return StreamingResponse(generator, media_type="text/event-stream")


@v1_router.get("/admin/strategist-reports", summary="View strategist reports")
async def v1_admin_strategist_reports(
    limit: int = QueryParam(default=20, le=100),
    _key: str = Depends(verify_admin_key),
):
    """View AI strategist nightly analysis reports (X-Admin-Key)."""
    try:
        from database import get_db_session
        from models import StrategistReport
        with get_db_session() as db:
            rows = db.query(StrategistReport).order_by(StrategistReport.created_at.desc()).limit(limit).all()
            return {"reports": [{
                "id": r.id,
                "report_text": r.report,
                "period_start": None,
                "period_end": None,
                "queries_analyzed": r.analyzed_count,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            } for r in rows], "total": len(rows)}
    except Exception as e:
        logger.error(f"Failed to list strategist reports: {e}")
        return {"reports": [], "total": 0}

@v1_router.post("/admin/strategist/trigger", summary="Manually trigger AI Strategist")
async def v1_admin_strategist_trigger(
    _key: str = Depends(verify_admin_key),
):
    """Manually trigger the AI Strategist to read the review queue and generate a report."""
    try:
        import asyncio
        strategist = Strategist()
        # Run in thread to not block event loop
        result = await asyncio.to_thread(strategist.analyze_low_confidence_queries)
        if result == "no_queries":
            raise HTTPException(status_code=400, detail="No pending queries in the review queue to analyze.")
        return {"status": "success", "message": "Strategist report generated successfully."}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to trigger strategist: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate report.")



class FeedbackRequest(BaseModel):
    query: str
    response: Optional[str] = None
    rating: int = Field(..., ge=1, le=5)
    comment: Optional[str] = None

    @field_validator("query", "response", "comment")
    @classmethod
    def sanitize_strings(cls, v: Optional[str]) -> Optional[str]:
        if v:
            return bleach.clean(v, tags=[], strip=True)
        return v

@v1_router.post("/query/feedback", summary="Submit user feedback")
async def v1_submit_feedback(
    body: FeedbackRequest,
    current_user: dict = Depends(get_current_user),
):
    """Submit feedback on a response (rating + optional comment)."""
    try:
        from database import get_db_session
        from models import Feedback
        with get_db_session() as db:
            db.add(Feedback(
                user_id=current_user["username"],
                query=body.query,
                response=body.response,
                rating=body.rating,
                comment=body.comment,
            ))
            db.commit()
    except Exception as e:
        logger.error(f"Feedback persistence failed: {e}")
    return {"message": "Feedback submitted", "rating": body.rating}


class GoldenAnswerRequest(BaseModel):
    query: str
    expected_answer: str
    relevant_chunk_ids: Optional[list[str]] = None
    tags: Optional[list[str]] = None

    @field_validator("query", "expected_answer")
    @classmethod
    def sanitize_strings(cls, v: str) -> str:
        return bleach.clean(v, tags=[], strip=True)

    @field_validator("tags")
    @classmethod
    def sanitize_tags(cls, v: Optional[list[str]]) -> Optional[list[str]]:
        if v:
            return [bleach.clean(tag, tags=[], strip=True) for tag in v]
        return v

@v1_router.post("/admin/golden-answers", summary="Add a golden answer")
async def v1_add_golden_answer(
    body: GoldenAnswerRequest,
    current_user: dict = Depends(require_permission("manage_users")),
):
    """Add a ground-truth golden answer for evaluation benchmarks."""
    try:
        from database import get_db_session
        from models import GoldenAnswer
        with get_db_session() as db:
            ga = GoldenAnswer(
                query=body.query,
                expected_answer=body.expected_answer,
                relevant_chunk_ids=body.relevant_chunk_ids,
                tags=body.tags,
            )
            db.add(ga)
            db.commit()
            return {"message": "Golden answer stored", "id": ga.id}
    except Exception as e:
        logger.error(f"Golden answer persistence failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to store golden answer")


@v1_router.get("/admin/golden-answers", summary="List golden answers")
async def v1_list_golden_answers(
    limit: int = QueryParam(default=50, le=500),
    current_user: dict = Depends(require_permission("manage_users")),
):
    """List all ground-truth golden answers."""
    try:
        from database import get_db_session
        from models import GoldenAnswer
        with get_db_session() as db:
            rows = db.query(GoldenAnswer).order_by(GoldenAnswer.created_at.desc()).limit(limit).all()
            return {"golden_answers": [{"id": r.id, "query": r.query, "expected_answer": r.expected_answer, "tags": r.tags} for r in rows], "total": len(rows)}
    except Exception as e:
        logger.error(f"Failed to list golden answers: {e}")
        return {"golden_answers": [], "total": 0}


# ───────────────────────────
# Stress Test Endpoints
# ───────────────────────────

class StressTestRequest(BaseModel):
    test_type: str = Field(default="all", pattern="^(all|bias|evasion|injection)$")
    queries: Optional[list[str]] = None


@v1_router.post("/stress-test/run", summary="Run stress test suite")
@limiter.limit("1/minute")
async def v1_stress_run(
    request: Request,
    body: StressTestRequest,
    _key: str = Depends(verify_admin_key),
):
    """Run adversarial stress tests via Server-Sent Events (X-Admin-Key, 1/min max)."""
    tester = StressTester()
    
    async def event_generator():
        if body.test_type == "bias":
            gen = tester.run_bias_test_stream()
        elif body.test_type == "evasion":
            gen = tester.run_evasion_test_stream()
        elif body.test_type == "injection":
            gen = tester.run_injection_test_stream()
        else:
            gen = tester.run_all_stream()
            
        async for event in gen:
            yield event
            
        yield "data: [DONE]\n\n"
        
        # Log to audit after completion
        audit.log(action="stress_test_executed", user_id="admin", metadata={
            "test_type": body.test_type,
            "streaming": True,
        })

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive"
        }
    )


# ───────────────────────────
# Evaluation Endpoints
# ───────────────────────────

@v1_router.get("/evaluations/summary", summary="Evaluation metrics summary")
async def v1_evaluations_summary():
    """Public summary of RAG evaluation metrics."""
    try:
        from database import get_db_session
        from models import Evaluation
        from sqlalchemy import func
        with get_db_session() as db:
            stats = db.query(
                func.avg(Evaluation.faithfulness).label("avg_faithfulness"),
                func.avg(Evaluation.relevance).label("avg_relevance"),
                func.avg(Evaluation.completeness).label("avg_completeness"),
                func.avg(Evaluation.f1_score).label("avg_f1"),
                func.avg(Evaluation.latency_ms).label("avg_latency_ms"),
                func.avg(Evaluation.estimated_cost_usd).label("avg_cost_usd"),
                func.count(Evaluation.id).label("total_queries"),
            ).first()
            # Queries by day (last 7 days)
            from datetime import timedelta
            cutoff = datetime.now(timezone.utc) - timedelta(days=7)
            daily = db.query(
                func.date(Evaluation.created_at).label("day"),
                func.count(Evaluation.id).label("count"),
            ).filter(
                Evaluation.created_at >= cutoff
            ).group_by(
                func.date(Evaluation.created_at)
            ).order_by("day").all()
            return {
                "avg_faithfulness": round(float(stats.avg_faithfulness or 0), 3),
                "avg_relevance": round(float(stats.avg_relevance or 0), 3),
                "avg_completeness": round(float(stats.avg_completeness or 0), 3),
                "avg_f1": round(float(stats.avg_f1 or 0), 3),
                "avg_latency_ms": round(float(stats.avg_latency_ms or 0), 1),
                "avg_cost_usd": round(float(stats.avg_cost_usd or 0), 6),
                "total_queries": int(stats.total_queries or 0),
                "queries_by_day": [{"day": str(d.day), "count": d.count} for d in daily],
            }
    except Exception as e:
        logger.error(f"Evaluations summary failed: {e}")
        return {"avg_faithfulness": 0, "avg_relevance": 0, "avg_completeness": 0, "avg_f1": 0, "avg_latency_ms": 0, "avg_cost_usd": 0, "total_queries": 0, "queries_by_day": []}


@v1_router.get("/evaluations/recent", summary="Recent evaluation records")
async def v1_evaluations_recent(
    limit: int = QueryParam(default=50, le=200),
    current_user: dict = Depends(get_current_user),
):
    """Last N evaluation records (JWT auth)."""
    try:
        from database import get_db_session
        from models import Evaluation
        with get_db_session() as db:
            rows = db.query(Evaluation).order_by(Evaluation.created_at.desc()).limit(limit).all()
            return {"evaluations": [{
                "id": r.id,
                "query": r.query,
                "faithfulness": r.faithfulness,
                "relevance": r.relevance,
                "completeness": r.completeness,
                "precision": r.precision,
                "recall": r.recall,
                "f1_score": r.f1_score,
                "confidence_score": r.confidence_score,
                "latency_ms": r.latency_ms,
                "estimated_cost_usd": r.estimated_cost_usd,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            } for r in rows], "total": len(rows)}
    except Exception as e:
        logger.error(f"Evaluations recent failed: {e}")
        return {"evaluations": [], "total": 0}


# ───────────────────────────
# Pipeline Status
# ───────────────────────────


@v1_router.get("/pipeline/status", summary="Live pipeline node status")
async def v1_pipeline_status(
    current_user: dict = Depends(get_current_user),
):
    """Returns in-memory dict of pipeline node statuses (JWT auth)."""
    return PIPELINE_STATUS


# ───────────────────────────
# SSE-Integrated Full Pipeline Chat
# ───────────────────────────

@v1_router.post("/chat/sse", summary="Full pipeline chat with SSE events")
async def v1_chat_sse(
    request: Request,
    body: ChatRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Full RAG pipeline with Server-Sent Events at each stage:
    planning → gatekeeper → retrieval → routing → agents → cost → evaluation → done
    """
    import asyncio
    import time
    from rag_pipeline import Planner, ToolExecution, ConditionalRouter as CondRouter
    from multi_agent import run_multi_agent, run_single_agent
    from auth import Gatekeeper as GK

    username = current_user["username"]
    user_role = current_user["role"]
    t_start = time.time()
    
    sse_queue = asyncio.Queue()

    async def pipeline_worker():
        try:
            update_pipeline_node("planner", "running")
            t_plan = time.time()
            await sse_queue.put({"event": "progress", "data": {"step": 1, "label": "Planning query", "icon": "brain"}})
            planner = Planner()
            plan = await planner.plan(body.question, sse_queue)
            update_pipeline_node("planner", "done", int((time.time() - t_plan)*1000))
            
            gk = GK()
            gate_result = gk.check_query(body.question, username)
            await sse_queue.put({"event": "gatekeeper", "data": {"status": gate_result["status"], "query": body.question[:100]}})
            
            if gate_result["status"] == "under_review":
                update_pipeline_node("human_validation", "running")
                await sse_queue.put({"event": "done", "data": {"answer": "Your query has been flagged for human review.", "status": "under_review"}})
                update_pipeline_node("human_validation", "done", 50)
                await sse_queue.put(None)
                return
                
            update_pipeline_node("data_sources", "running")
            update_pipeline_node("tool_execution", "running")
            t_ret = time.time()
            await sse_queue.put({"event": "progress", "data": {"step": 2, "label": "Retrieving context", "icon": "search"}})
            tool_exec = ToolExecution(vector_store, sse_queue)
            chunks = await tool_exec.execute(plan, body.question)
            ret_dur = int((time.time() - t_ret)*1000)
            update_pipeline_node("tool_execution", "done", ret_dur)
            update_pipeline_node("data_sources", "done", ret_dur)
            
            update_pipeline_node("router", "running")
            t_route = time.time()
            await sse_queue.put({"event": "progress", "data": {"step": 3, "label": "Routing query", "icon": "route"}})
            router = CondRouter(sse_queue)
            route_decision = await router.route(chunks)
            update_pipeline_node("router", "done", int((time.time() - t_route)*1000))
            
            await sse_queue.put({"event": "progress", "data": {"step": 4, "label": "Running generation agents", "icon": "bot"}})
            decision = route_decision["decision"]
            final_response = ""
            
            t_agent = time.time()
            if decision == "multi_agent" or decision == "multi-agent":
                update_pipeline_node("multi_agent", "running")
                state = await run_multi_agent(body.question, chunks, {"user_role": user_role, "format": body.format}, sse_queue)
                final_response = state["final_response"]
                update_pipeline_node("multi_agent", "done", int((time.time() - t_agent)*1000))
            elif decision == "human_validation" or decision == "human review":
                update_pipeline_node("human_validation", "running")
                final_response = "This query requires human review due to low confidence."
                await sse_queue.put({"event": "done", "data": {"answer": final_response, "status": "under_review"}})
                update_pipeline_node("human_validation", "done", 50)
                await sse_queue.put(None)
                return
            else:
                update_pipeline_node("agent_1", "running")
                state = await run_single_agent(body.question, chunks, sse_queue, metadata={"user_role": user_role, "format": body.format})
                final_response = state["final_response"]
                update_pipeline_node("agent_1", "done", int((time.time() - t_agent)*1000))
                
            # --- Disabled evaluation to maximize speed ---
            # try:
            #     evaluator = Evaluator()
            #     eval_result = await evaluator.evaluate_async(...)
            # except Exception as e:
            #     logger.error(f"Evaluation failed: {e}")
                
            total_ms = int((time.time() - t_start) * 1000)
            _track_analytics("sse_pipeline_complete", user_id=username, metadata={
                "strategy": plan.get("strategy"),
                "agent_path": decision,
                "total_ms": total_ms,
            })
            
            # Persist the chat to database
            await _persist_chat(username, body.question, final_response)
            
            await sse_queue.put({"event": "done", "data": {
                "answer": final_response,
                "confidence": route_decision["confidence"],
                "strategy": plan.get("strategy"),
                "agent_path": decision,
                "total_pipeline_ms": total_ms
            }})
            
        except Exception as e:
            logger.error(f"Pipeline processing failed: {e}", exc_info=True)
            await sse_queue.put({"event": "error", "data": {"message": str(e)}})
        finally:
            await sse_queue.put(None) # Signal termination

    async def pipeline_generator():
        # Start processing task
        asyncio.create_task(pipeline_worker())
        
        while True:
            msg = await sse_queue.get()
            if msg is None:
                break
            yield sse_manager._format_sse(msg["event"], msg["data"])

    return StreamingResponse(
        pipeline_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ───────────────────────────
# SSE Streaming Endpoints
# ───────────────────────────

@v1_router.get("/events/stream", summary="SSE event stream")
async def v1_sse_stream(
    channel: str = QueryParam(default="default"),
    current_user: dict = Depends(get_current_user),
):
    """Subscribe to real-time Server-Sent Events with automatic heartbeat."""
    queue = sse_manager.subscribe(channel)

    async def event_generator():
        async for message in sse_manager.stream(queue, channel):
            yield message

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@v1_router.post("/events/publish", summary="Publish an SSE event")
async def v1_sse_publish(
    event: str = "update",
    channel: str = "default",
    current_user: dict = Depends(require_permission("execute")),
):
    """Publish an event to all SSE subscribers on a channel."""
    await sse_manager.publish(event, {
        "message": f"Event from {current_user['username']}",
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }, channel)
    return {"published": True, "channel": channel, "subscribers": sse_manager.active_connections(channel)}


@v1_router.get("/events/status", summary="SSE connection stats")
async def v1_sse_status(
    current_user: dict = Depends(get_current_user),
):
    """View active SSE subscriber count per channel."""
    return {
        "channels": {ch: len(subs) for ch, subs in sse_manager._channels.items()},
        "total_subscribers": sum(len(s) for s in sse_manager._channels.values()),
    }


# ── Mount V1 Router ──
app.include_router(v1_router)


# ═══════════════════════════════════════════
# Entry Point
# ═══════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info",
    )